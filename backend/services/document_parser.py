"""
Document Parser — Stage 1
==========================
Extracts structured text from PDF, DOCX, and TXT files.

Returns List[PageBlock] preserving reading order, headings, and page numbers.
Also extracts a defined-terms glossary for chunk context injection.

Design decisions:
- PyMuPDF (fitz) for PDFs: layout-aware, preserves reading order in multi-column docs
- python-docx for DOCX: style-based heading detection is more reliable than font size
- Both return the same List[PageBlock] interface for pipeline compatibility
- No OCR: scanned PDFs are a known limitation (documented in README)
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import structlog
from docx import Document
from docx.oxml.ns import qn

logger = structlog.get_logger(__name__)


@dataclass
class PageBlock:
    """A single block of text from a parsed document."""
    page_num: int
    text: str
    is_heading: bool
    font_size: float
    bbox: tuple[float, float, float, float] = field(default=(0, 0, 0, 0))  # x0,y0,x1,y1

    @property
    def text_stripped(self) -> str:
        return self.text.strip()

    def is_empty(self) -> bool:
        return not self.text.strip()


@dataclass
class ParsedDocument:
    """Output of document parsing stage."""
    blocks: list[PageBlock]
    defined_terms: dict[str, str]  # term → definition
    total_pages: int
    raw_text: str  # Full concatenated text for classification
    file_type: str  # "pdf" | "docx" | "txt"


# ── PDF Parsing ───────────────────────────────────────────────────────────────

def _estimate_body_font_size(page: fitz.Page) -> float:
    """
    Determine the modal (most common) font size on a page — used as body text baseline.
    Text significantly larger than this is likely a heading.
    """
    font_sizes = []
    for block in page.get_text("dict")["blocks"]:
        if block["type"] != 0:  # 0 = text block
            continue
        for line in block["lines"]:
            for span in line["spans"]:
                size = round(span["size"])
                if size > 0:
                    font_sizes.append(size)

    if not font_sizes:
        return 12.0

    # Return modal font size
    return max(set(font_sizes), key=font_sizes.count)


def parse_pdf(file_path: str) -> ParsedDocument:
    """
    Parse a PDF file using PyMuPDF with layout-preserving block extraction.

    Approach:
    1. Extract text blocks per page using get_text("dict") for structure
    2. Sort blocks by (page, y0, x0) to restore reading order
    3. Classify blocks as headings based on font size vs body baseline
    4. Filter out headers/footers (repeated text at page extremes)
    """
    doc = fitz.open(file_path)
    all_blocks: list[PageBlock] = []

    # Collect body font sizes across first 5 pages for stable baseline
    body_sizes = []
    for page_num in range(min(5, len(doc))):
        body_sizes.append(_estimate_body_font_size(doc[page_num]))
    body_font_size = max(set(body_sizes), key=body_sizes.count) if body_sizes else 12.0
    heading_threshold = body_font_size * 1.15  # 15% larger = likely heading

    # Track repeated short strings (headers/footers) to filter them
    short_string_counts: dict[str, int] = {}

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_dict = page.get_text("dict")
        page_height = page.rect.height

        raw_blocks = []
        for block in page_dict["blocks"]:
            if block["type"] != 0:
                continue

            # Aggregate all spans in the block into one text string
            block_text = ""
            max_font_size = 0.0
            is_bold = False

            for line in block["lines"]:
                line_text = ""
                for span in line["spans"]:
                    line_text += span["text"]
                    if span["size"] > max_font_size:
                        max_font_size = span["size"]
                    # Bold detection via font flags (bit 4 = bold)
                    if span["flags"] & 16:
                        is_bold = True
                block_text += line_text + "\n"

            block_text = block_text.strip()
            if not block_text:
                continue

            # Track short strings for header/footer detection
            if len(block_text) < 80:
                short_string_counts[block_text] = short_string_counts.get(block_text, 0) + 1

            bbox = block["bbox"]
            raw_blocks.append({
                "text": block_text,
                "font_size": max_font_size,
                "is_bold": is_bold,
                "bbox": bbox,
                "y0": bbox[1],
                "x0": bbox[0],
            })

        # Sort by y0 (top to bottom) then x0 (left to right) for reading order
        raw_blocks.sort(key=lambda b: (b["y0"], b["x0"]))

        for block_data in raw_blocks:
            text = block_data["text"]
            font_size = block_data["font_size"]

            # Classify as heading: larger than body font, bold, or ALL CAPS short text
            is_heading = (
                font_size >= heading_threshold
                or block_data["is_bold"]
                or (text.isupper() and len(text) < 100)
            )

            # Filter out obvious page numbers (just a number)
            if re.match(r"^\d+$", text.strip()):
                continue

            all_blocks.append(PageBlock(
                page_num=page_num + 1,  # 1-indexed
                text=text,
                is_heading=is_heading,
                font_size=font_size,
                bbox=block_data["bbox"],
            ))

    # Remove headers/footers (strings appearing on 3+ pages)
    repeated = {text for text, count in short_string_counts.items() if count >= 3}
    filtered_blocks = [b for b in all_blocks if b.text not in repeated]

    raw_text = "\n".join(b.text for b in filtered_blocks)
    defined_terms = extract_defined_terms(raw_text)

    logger.info(
        "pdf_parsed",
        file=file_path,
        pages=len(doc),
        blocks=len(filtered_blocks),
        defined_terms=len(defined_terms),
    )

    doc.close()
    return ParsedDocument(
        blocks=filtered_blocks,
        defined_terms=defined_terms,
        total_pages=len(doc),
        raw_text=raw_text,
        file_type="pdf",
    )


# ── DOCX Parsing ──────────────────────────────────────────────────────────────

def parse_docx(file_path: str) -> ParsedDocument:
    """
    Parse a DOCX file using python-docx.

    Heading detection uses paragraph.style.name (Heading 1, 2, 3) which is
    more reliable than font-size heuristics in Word documents.

    Tables are included because payment schedules and term sheets use them.
    """
    doc = Document(file_path)
    all_blocks: list[PageBlock] = []
    page_num = 1  # DOCX has no native page concept; we approximate

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        is_heading = para.style.name.startswith("Heading")

        # Approximate page numbers by counting words (rough: 250 words/page)
        word_count = len(text.split())

        all_blocks.append(PageBlock(
            page_num=page_num,
            text=text,
            is_heading=is_heading,
            font_size=14.0 if is_heading else 11.0,
        ))

    # Extract tables
    for table in doc.tables:
        table_text_parts = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_text_parts.append(row_text)
        if table_text_parts:
            table_text = "\n".join(table_text_parts)
            all_blocks.append(PageBlock(
                page_num=page_num,
                text=table_text,
                is_heading=False,
                font_size=11.0,
            ))

    raw_text = "\n".join(b.text for b in all_blocks)
    defined_terms = extract_defined_terms(raw_text)

    # Count approximate pages from word count
    total_words = len(raw_text.split())
    total_pages = max(1, total_words // 250)

    logger.info(
        "docx_parsed",
        file=file_path,
        paragraphs=len(all_blocks),
        defined_terms=len(defined_terms),
        approx_pages=total_pages,
    )

    return ParsedDocument(
        blocks=all_blocks,
        defined_terms=defined_terms,
        total_pages=total_pages,
        raw_text=raw_text,
        file_type="docx",
    )


# ── TXT Parsing ───────────────────────────────────────────────────────────────

def parse_txt(file_path: str) -> ParsedDocument:
    """Plain text parsing with heuristic heading detection."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        raw_text = f.read()

    lines = raw_text.split("\n")
    all_blocks: list[PageBlock] = []
    page_num = 1

    for line in lines:
        text = line.strip()
        if not text:
            continue

        # Heuristic: ALL CAPS lines or numbered section headers are headings
        is_heading = (
            text.isupper()
            or bool(re.match(r"^(SECTION|ARTICLE|\d+\.|[IVXLC]+\.)\s+[A-Z]", text))
        )

        all_blocks.append(PageBlock(
            page_num=page_num,
            text=text,
            is_heading=is_heading,
            font_size=14.0 if is_heading else 11.0,
        ))

    defined_terms = extract_defined_terms(raw_text)
    total_pages = max(1, len(raw_text.split()) // 250)

    return ParsedDocument(
        blocks=all_blocks,
        defined_terms=defined_terms,
        total_pages=total_pages,
        raw_text=raw_text,
        file_type="txt",
    )


# ── Defined Term Extraction ────────────────────────────────────────────────────

def extract_defined_terms(text: str) -> dict[str, str]:
    """
    Extract a glossary of defined terms from the full contract text.

    Patterns recognized:
    1. "Term" means ... (quoted term with definition)
    2. "Term" shall mean ...
    3. (as used herein, "Term" means ...)
    4. "Term" has the meaning given in Section X
    Also: capitalized phrases used 3+ times (CamelCase or ALL_CAPS compound terms)
    """
    defined_terms: dict[str, str] = {}

    # Pattern 1 & 2: "X" means/shall mean ...
    pattern_explicit = re.compile(
        r'"([A-Z][^"]{1,60})"\s+(?:shall\s+)?means?\s+([^.;]{10,300})[.;]',
        re.IGNORECASE,
    )
    for match in pattern_explicit.finditer(text):
        term = match.group(1).strip()
        definition = match.group(2).strip()
        if len(term) > 2:
            defined_terms[term] = definition[:200]

    # Pattern 3: ("Term" means ...)
    pattern_paren = re.compile(
        r'\("([A-Z][^"]{1,60})"\s+(?:shall\s+)?means?\s+([^)]{10,300})\)',
        re.IGNORECASE,
    )
    for match in pattern_paren.finditer(text):
        term = match.group(1).strip()
        definition = match.group(2).strip()
        if term not in defined_terms:
            defined_terms[term] = definition[:200]

    # Capitalized compound terms used 3+ times (heuristic for common legal shorthand)
    camel_pattern = re.compile(r'\b([A-Z][a-z]+(?:[A-Z][a-z]+)+)\b')
    camel_counts: dict[str, int] = {}
    for match in camel_pattern.finditer(text):
        term = match.group(1)
        if term not in ("This", "The", "Such", "Each", "Any"):
            camel_counts[term] = camel_counts.get(term, 0) + 1

    for term, count in camel_counts.items():
        if count >= 3 and term not in defined_terms:
            defined_terms[term] = f"[Recurring term used {count} times — see contract for definition]"

    return defined_terms


# ── Main Entry Point ──────────────────────────────────────────────────────────

def parse_document(file_path: str) -> ParsedDocument:
    """
    Route to the correct parser based on file extension.
    Called by the contracts router after upload validation.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".txt": parse_txt,
    }

    if ext not in parsers:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(parsers.keys())}")

    logger.info("parsing_document", file=file_path, type=ext)
    return parsers[ext](file_path)
