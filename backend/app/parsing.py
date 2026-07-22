"""
Document parsing: PDF (PyMuPDF, OCR fallback) and DOCX (python-docx).

Honest limitation (see audit): multi-column layouts and embedded tables WILL
produce jumbled text. This is not silently hidden — callers get
`layout_warning=True` when the heuristic below suspects a multi-column PDF,
so the frontend can tell the user to double check results.
"""
import io
import logging

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.config import settings

logger = logging.getLogger("clauseguard.parsing")

# Tesseract is an OS-level binary. It may not be installed in every
# deployment environment. Feature-detect instead of assuming it's there —
# a naive `import pytesseract` + call will crash the whole worker if the
# binary is missing, taking down every document in the queue with it.
_OCR_AVAILABLE = False
if settings.OCR_ENABLED:
    try:
        import pytesseract
        from PIL import Image
        pytesseract.get_tesseract_version()
        _OCR_AVAILABLE = True
    except Exception as e:  # binary missing, or pytesseract/Pillow not installed
        logger.warning(f"OCR unavailable, falling back to text-only extraction: {e}")
        _OCR_AVAILABLE = False


class ParseError(Exception):
    def __init__(self, code: str, message: str):
        self.code = code
        self.message = message
        super().__init__(message)


class ParseResult:
    def __init__(self, text: str, page_count: int, used_ocr: bool, layout_warning: bool = False):
        self.text = text
        self.page_count = page_count
        self.used_ocr = used_ocr
        self.layout_warning = layout_warning


def _detect_multicolumn(page: "fitz.Page") -> bool:
    """
    Heuristic: if text blocks on a page cluster into two clearly separated
    horizontal bands (left half / right half), it's very likely a two-column
    layout, which PyMuPDF's raw reading order will scramble. This is a
    warning signal, not a fix — actual column-aware extraction is out of
    scope for v1 (documented in the original spec, section 13.1).
    """
    try:
        blocks = page.get_text("blocks")
        if len(blocks) < 6:
            return False
        page_width = page.rect.width
        midpoint = page_width / 2
        left_blocks = sum(1 for b in blocks if b[2] < midpoint + 5)
        right_blocks = sum(1 for b in blocks if b[0] > midpoint - 5)
        total = len(blocks)
        return (left_blocks / total > 0.3) and (right_blocks / total > 0.3)
    except Exception:
        return False


def extract_text_from_pdf(file_bytes: bytes) -> ParseResult:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        raise ParseError("parse_failed", f"Could not open PDF: {e}")

    if doc.needs_pass:
        raise ParseError("parse_failed", "This PDF is password-protected. Remove the password and try again.")

    page_count = len(doc)
    if page_count == 0:
        raise ParseError("parse_failed", "PDF has no pages.")
    if page_count > settings.MAX_PAGES:
        raise ParseError("too_large", f"Document has {page_count} pages; limit is {settings.MAX_PAGES}.")

    pages_text = []
    layout_warning = False
    for page in doc:
        pages_text.append(page.get_text())
        if _detect_multicolumn(page):
            layout_warning = True

    full_text = "\n".join(pages_text)
    avg_chars_per_page = len(full_text.strip()) / max(page_count, 1)

    used_ocr = False
    if avg_chars_per_page < settings.OCR_CHAR_THRESHOLD_PER_PAGE:
        if not _OCR_AVAILABLE:
            raise ParseError(
                "ocr_unavailable",
                "This looks like a scanned/image-based PDF, and OCR is not available on this "
                "server (tesseract binary not installed). Install it with "
                "`apt-get install tesseract-ocr` or upload a text-based PDF instead.",
            )
        full_text = _extract_via_ocr(doc)
        used_ocr = True

    doc.close()
    return ParseResult(text=full_text, page_count=page_count, used_ocr=used_ocr, layout_warning=layout_warning)


def _extract_via_ocr(doc: "fitz.Document") -> str:
    all_text = []
    for page in doc:
        mat = fitz.Matrix(2.0, 2.0)
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        text = pytesseract.image_to_string(img, lang="eng")
        all_text.append(text)
    return "\n".join(all_text)


def extract_text_from_docx(file_bytes: bytes) -> ParseResult:
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ParseError("parse_failed", f"Could not open DOCX: {e}")

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables in DOCX are not walked by doc.paragraphs — pull them separately
    # so at least the raw cell text isn't silently dropped (still won't
    # preserve table structure, same limitation as the PDF path).
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                table_text.append(" | ".join(cells))

    full_text = "\n".join(paragraphs + table_text)
    word_count = len(full_text.split())
    estimated_pages = max(1, word_count // 350)

    return ParseResult(text=full_text, page_count=estimated_pages, used_ocr=False)


def parse_document(file_bytes: bytes, file_type: str) -> ParseResult:
    if len(file_bytes) == 0:
        raise ParseError("parse_failed", "File is empty.")

    if file_type == "pdf":
        result = extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        result = extract_text_from_docx(file_bytes)
    else:
        raise ParseError("unsupported_type", f"Unsupported file type: {file_type}")

    word_count = len(result.text.split())
    if word_count < settings.MIN_EXTRACTED_WORDS:
        raise ParseError(
            "text_too_short",
            f"Only extracted {word_count} words. This may be a scanned image without "
            f"OCR-compatible text, or not a real contract.",
        )

    return result
